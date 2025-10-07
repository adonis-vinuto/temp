import io
import re
import tempfile
from pathlib import Path
from typing import List, Tuple, Dict, Any

import pdfplumber
import pandas as pd
import camelot
import pypdfium2 as pdfium
from docx import Document

from ...domain.extract_text_schema import FileSchema, PageSchema
from ...domain.usage_schema import UsageSchema
from .text_refiner import text_resume, text_file_name
from ...application.ocr_service import ocr_service
from ..simple_usage_callback import SimpleUsageCallback


def detect_file_type(file_name: str) -> str:
    """Detecta o tipo do arquivo pela extensão."""
    suffix = Path(file_name).suffix.lower()

    mapping = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".xlsx": "excel",
        ".xls": "excel",
        ".xlsm": "excel",
        ".csv": "excel",
    }

    return mapping.get(suffix, "unknown")


def is_tabular_text(text: str) -> bool:
    """Verifica se o texto de uma página tem características tabulares."""
    num_ratio = len(re.findall(r"\d+[\.,]\d+", text)) / (len(text.split()) + 1)

    has_keywords = bool(re.search(r"M[êe]s\s*\d|Total|Receita|Custos", text, re.I))

    short_lines = sum(1 for l in text.splitlines() if len(l) < 40)

    return (num_ratio > 0.15 and has_keywords) or short_lines > 20


def split_text_smart(text: str, max_length: int = 1000) -> List[str]:
    """Divide o texto sem quebrar frases."""
    parts, start = [], 0

    while start < len(text):
        end = start + max_length

        if end >= len(text):
            parts.append(text[start:].strip())
            break

        split_pos = max(
            text.rfind(".", start, end),
            text.rfind("\n", start, end),
            text.rfind(" ", start, end),
        )

        if split_pos <= start:
            split_pos = end

        parts.append(text[start:split_pos].strip())

        start = split_pos

    return parts


def extract_pdf_pages_from_bytes(pdf_bytes: bytes, file_name: str) -> Tuple[FileSchema, Dict[str, Any]]:
    """
    Extrai o conteúdo de um PDF (texto, tabelas e OCR se necessário).
    Retorna tupla (FileSchema, usage_dict)
    """
    pages: List[PageSchema] = []
    text_found = False
    usage_callback = SimpleUsageCallback()

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()

                if text and text.strip():
                    text_found = True

                    if is_tabular_text(text):
                        pages.extend(_extract_pdf_table_page(pdf_bytes, i))

                    else:
                        for part in split_text_smart(text):
                            pages.append(PageSchema(page_number=i, text=part))

                else:
                    # Fallback OCR
                    pdf_page = pdfium.PdfDocument(pdf_bytes).get_page(i - 1)

                    ocr_text = ocr_service(pdf_page)

                    if ocr_text:
                        text_found = True

                        for part in split_text_smart(ocr_text):
                            pages.append(PageSchema(page_number=i, text=part))

    except Exception as e:
        print(f"[ERRO] pdfplumber falhou: {e}. Revertendo para OCR completo.")

        pdf_doc = pdfium.PdfDocument(pdf_bytes)

        for i in range(len(pdf_doc)):
            ocr_text = ocr_service(pdf_doc.get_page(i))

            if ocr_text:
                text_found = True

                for part in split_text_smart(ocr_text):
                    pages.append(PageSchema(page_number=i + 1, text=part))

    if not text_found:
        raise ValueError("Nenhum conteúdo extraível encontrado, mesmo com OCR.")

    # Usar o mesmo callback para ambas as chamadas
    file_name_result = text_file_name(pages, usage_callback)
    resume_result = text_resume(pages, usage_callback)

    return FileSchema(
        file_name=file_name_result,
        file_type="pdf",
        resume=resume_result,
        pages=pages,
    ), usage_callback.get_usage_dict()


def _extract_pdf_table_page(pdf_bytes: bytes, page_number: int) -> List[PageSchema]:
    """Extrai tabelas de uma página PDF usando Camelot."""
    pages = []

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(pdf_bytes)

        tmp_path = tmp_file.name

    try:
        tables = camelot.read_pdf(tmp_path, pages=str(page_number), flavor="stream")

        for table in tables:
            df = table.df.copy()

            text = df.to_string(index=False)

            pages.append(PageSchema(page_number=page_number, text=text.strip()))

    finally:
        from os import remove

        remove(tmp_path)

    return pages


def extract_docx_content_from_bytes(docx_bytes: bytes, file_name: str) -> Tuple[FileSchema, Dict[str, Any]]:
    """Extrai texto de arquivos .docx. Retorna tupla (FileSchema, usage_dict)"""
    doc = Document(io.BytesIO(docx_bytes))
    usage_callback = SimpleUsageCallback()

    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    parts = split_text_smart(text)

    pages = [PageSchema(page_number=i, text=part) for i, part in enumerate(parts, start=1)]

    # Usar o mesmo callback para ambas as chamadas
    file_name_result = text_file_name(pages, usage_callback)
    resume_result = text_resume(pages, usage_callback)

    return FileSchema(
        file_name=file_name_result,
        file_type="docx",
        resume=resume_result,
        pages=pages,
    ), usage_callback.get_usage_dict()


def extract_txt_content_from_bytes(txt_bytes: bytes, file_name: str) -> Tuple[FileSchema, Dict[str, Any]]:
    """Extrai texto de arquivos .txt. Retorna tupla (FileSchema, usage_dict)"""
    content = txt_bytes.decode("utf-8", errors="ignore")
    usage_callback = SimpleUsageCallback()

    parts = split_text_smart(content)

    pages = [PageSchema(page_number=i, text=part) for i, part in enumerate(parts, start=1)]

    # Usar o mesmo callback para ambas as chamadas
    file_name_result = text_file_name(pages, usage_callback)
    resume_result = text_resume(pages, usage_callback)

    return FileSchema(
        file_name=file_name_result,
        file_type="txt",
        resume=resume_result,
        pages=pages,
    ), usage_callback.get_usage_dict()


def extract_excel_content_from_bytes(excel_bytes: bytes, file_name: str) -> Tuple[FileSchema, Dict[str, Any]]:
    """Extrai conteúdo de planilhas Excel ou CSV. Retorna tupla (FileSchema, usage_dict)"""
    pages = []
    usage_callback = SimpleUsageCallback()

    try:
        dfs = (
            {"Sheet1": pd.read_csv(io.BytesIO(excel_bytes))}
            if Path(file_name).suffix.lower() == ".csv"
            else pd.read_excel(io.BytesIO(excel_bytes), sheet_name=None)
        )

        for idx, (sheet, df) in enumerate(dfs.items(), start=1):
            df = df.dropna(how="all").dropna(axis=1, how="all").fillna("")

            df.columns = [col if not str(col).startswith("Unnamed") else "" for col in df.columns]

            if df.empty:
                continue

            try:
                table_text = f"# Planilha: {sheet}\n\n{df.to_markdown(index=False)}"

            except Exception:
                table_text = f"# Planilha: {sheet}\n\n{df.to_string(index=False)}"

            for part in split_text_smart(table_text, max_length=3000):
                pages.append(PageSchema(page_number=idx, text=part.strip()))

        # Usar o mesmo callback para ambas as chamadas
        file_name_result = text_file_name(pages, usage_callback)
        resume_result = text_resume(pages, usage_callback)

        return FileSchema(
            file_name=file_name_result,
            file_type="excel",
            resume=resume_result,
            pages=pages,
        ), usage_callback.get_usage_dict()

    except Exception as e:
        raise ValueError(f"Erro ao processar Excel: {e}")


def extract_file_content_from_bytes(file_bytes: bytes, file_name: str) -> Tuple[FileSchema, UsageSchema]:
    """
    Função principal — escolhe o extrator com base no tipo de arquivo.
    Retorna tupla (FileSchema, usage_dict)
    """
    file_type = detect_file_type(file_name)

    extractors = {
        "pdf": extract_pdf_pages_from_bytes,
        "docx": extract_docx_content_from_bytes,
        "txt": extract_txt_content_from_bytes,
        "excel": extract_excel_content_from_bytes,
    }

    if file_type not in extractors:
        raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
    
    return extractors[file_type](file_bytes, file_name)
